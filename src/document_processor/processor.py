from typing import List, Dict, Any, Optional, Tuple
from sqlalchemy.orm import Session
import os
import asyncio
import uuid

from src.database import Document, DocumentChunk, DocumentStatus, DocumentType, DocumentSource
from src.llm_connector.factory import get_embedding_model


async def process_document(
    document_id: str,
    db: Session,
) -> bool:
    """
    Process a document and create chunks.
    
    Args:
        document_id: The ID of the document to process
        db: Database session
        
    Returns:
        True if processing was successful, False otherwise
    """
    # Get the document
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        return False
    
    try:
        # Update document status to processing
        document.status = DocumentStatus.PROCESSING
        db.commit()
        
        # Extract text based on document type
        text = await extract_text(document)
        
        # Chunk the text using document type-specific strategy
        chunks = chunk_text(text, doc_type=document.doc_type)
        
        # Get embedding model
        embedding_model = get_embedding_model(db)
        
        # Process each chunk
        for i, chunk_text in enumerate(chunks):
            # Create chunk
            chunk = DocumentChunk(
                document_id=document_id,
                content=chunk_text,
                chunk_index=i,
                meta_data={"position": i, "total_chunks": len(chunks)},
            )
            db.add(chunk)
            db.commit()
            db.refresh(chunk)
            
            # Generate embedding
            embedding = await embedding_model.get_embedding(chunk_text)
            
            # Save embedding to file
            embedding_dir = os.path.join("data", "embeddings", document_id)
            os.makedirs(embedding_dir, exist_ok=True)
            embedding_file = os.path.join(embedding_dir, f"{chunk.id}.npy")
            
            # In a real implementation, we would save the embedding to a file
            # For now, just update the embedding_file field
            chunk.embedding_file = embedding_file
            db.commit()
        
        # Update document status to completed
        document.status = DocumentStatus.COMPLETED
        db.commit()
        
        return True
    
    except Exception as e:
        # Update document status to failed
        document.status = DocumentStatus.FAILED
        document.meta_data = document.meta_data or {}
        document.meta_data["error"] = str(e)
        db.commit()
        
        return False


import pypdf
import docx
import markdown
from docutils.core import publish_string
from bs4 import BeautifulSoup
import io
import os

async def extract_text(document: Document) -> str:
    """
    Extract text from a document.
    
    Args:
        document: The document to extract text from
        
    Returns:
        The extracted text
    """
    # Check if file exists for file-based documents
    if document.source in [DocumentSource.UPLOAD, DocumentSource.GITHUB] and document.file_path:
        if not os.path.exists(document.file_path):
            raise FileNotFoundError(f"Document file not found: {document.file_path}")
    
    try:
        if document.doc_type == DocumentType.PDF:
            return await extract_text_from_pdf(document)
        
        elif document.doc_type == DocumentType.DOCX:
            return await extract_text_from_docx(document)
        
        elif document.doc_type == DocumentType.MARKDOWN:
            return await extract_text_from_markdown(document)
        
        elif document.doc_type == DocumentType.RST:
            return await extract_text_from_rst(document)
        
        elif document.doc_type == DocumentType.TEXT:
            return await extract_text_from_text(document)
        
        elif document.doc_type == DocumentType.HTML:
            return await extract_text_from_html(document)
        
        elif document.doc_type == DocumentType.FORM:
            # Form data is already in text format in the metadata
            return await extract_text_from_form(document)
        
        else:
            return f"Unknown document type: {document.doc_type}. Unable to extract text."
    
    except Exception as e:
        # Log the error and return a message
        print(f"Error extracting text from {document.title}: {str(e)}")
        return f"Error extracting text: {str(e)}"


async def extract_text_from_pdf(document: Document) -> str:
    """Extract text from a PDF document."""
    # Use pypdf to extract text
    text = ""
    
    try:
        with open(document.file_path, "rb") as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Extract metadata if available
            metadata = pdf_reader.metadata
            if metadata:
                text += f"Title: {metadata.title or document.title}\n"
                if metadata.author:
                    text += f"Author: {metadata.author}\n"
                if metadata.subject:
                    text += f"Subject: {metadata.subject}\n"
                if metadata.creator:
                    text += f"Creator: {metadata.creator}\n"
                text += "\n"
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
    
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    return text


async def extract_text_from_docx(document: Document) -> str:
    """Extract text from a DOCX document."""
    # Use python-docx to extract text
    text = ""
    
    try:
        doc = docx.Document(document.file_path)
        
        # Extract document properties if available
        core_properties = doc.core_properties
        if core_properties:
            if core_properties.title:
                text += f"Title: {core_properties.title}\n"
            if core_properties.author:
                text += f"Author: {core_properties.author}\n"
            if core_properties.subject:
                text += f"Subject: {core_properties.subject}\n"
            text += "\n"
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += " | ".join(row_text) + "\n"
            text += "\n"
    
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    return text


async def extract_text_from_markdown(document: Document) -> str:
    """Extract text from a Markdown document."""
    try:
        with open(document.file_path, "r", encoding="utf-8") as file:
            md_content = file.read()
        
        # Return the raw markdown content
        # We could convert to HTML and then extract text, but for RAG purposes,
        # keeping the original markdown formatting is often better
        return md_content
    
    except Exception as e:
        raise Exception(f"Error extracting text from Markdown: {str(e)}")


async def extract_text_from_rst(document: Document) -> str:
    """Extract text from an RST document."""
    try:
        with open(document.file_path, "r", encoding="utf-8") as file:
            rst_content = file.read()
        
        # Option 1: Return the raw RST content
        # Similar to markdown, keeping the original formatting can be useful for RAG
        return rst_content
        
        # Option 2: Convert to HTML and extract text
        # html = publish_string(rst_content, writer_name='html').decode('utf-8')
        # soup = BeautifulSoup(html, 'lxml')
        # return soup.get_text(separator='\n')
    
    except Exception as e:
        raise Exception(f"Error extracting text from RST: {str(e)}")


async def extract_text_from_text(document: Document) -> str:
    """Extract text from a plain text document."""
    try:
        with open(document.file_path, "r", encoding="utf-8") as file:
            return file.read()
    
    except Exception as e:
        raise Exception(f"Error extracting text from text file: {str(e)}")


async def extract_text_from_html(document: Document) -> str:
    """Extract text from an HTML document."""
    try:
        # For file-based HTML
        if document.file_path:
            with open(document.file_path, "r", encoding="utf-8") as file:
                html_content = file.read()
        # For URL-based HTML
        elif document.url:
            import httpx
            async with httpx.AsyncClient() as client:
                response = await client.get(document.url)
                response.raise_for_status()
                html_content = response.text
        else:
            raise ValueError("No file path or URL provided for HTML document")
        
        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, 'lxml')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Extract text
        text = soup.get_text(separator='\n')
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    except Exception as e:
        raise Exception(f"Error extracting text from HTML: {str(e)}")


async def extract_text_from_form(document: Document) -> str:
    """Extract text from form data."""
    try:
        # Form data should be stored in the metadata
        if not document.meta_data:
            return "No form data found."
        
        # Convert form data to text
        text = f"Form: {document.title}\n\n"
        
        for key, value in document.meta_data.items():
            if key != "error":  # Skip error messages
                text += f"{key}: {value}\n"
        
        return text
    
    except Exception as e:
        raise Exception(f"Error extracting text from form data: {str(e)}")


def chunk_text(text: str, doc_type: DocumentType = None, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Chunk text into smaller pieces using an appropriate strategy based on document type.
    
    Args:
        text: The text to chunk
        doc_type: The document type (optional)
        chunk_size: The maximum size of each chunk
        overlap: The overlap between chunks
        
    Returns:
        A list of text chunks
    """
    # If text is shorter than chunk_size, return it as a single chunk
    if len(text) <= chunk_size:
        return [text]
    
    # Choose chunking strategy based on document type
    if doc_type == DocumentType.MARKDOWN or doc_type == DocumentType.RST:
        return chunk_by_heading(text, chunk_size, overlap)
    elif doc_type == DocumentType.PDF:
        return chunk_by_page_then_paragraph(text, chunk_size, overlap)
    else:
        # Default to paragraph chunking for other document types
        return chunk_by_paragraph(text, chunk_size, overlap)


def chunk_by_paragraph(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Chunk text by paragraphs."""
    # Split text into paragraphs
    paragraphs = text.split("\n\n")
    
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        # If paragraph itself is longer than chunk_size, split it by sentences
        if len(paragraph) > chunk_size:
            sentence_chunks = chunk_by_sentence(paragraph, chunk_size, overlap)
            for sentence_chunk in sentence_chunks:
                if len(current_chunk) + len(sentence_chunk) > chunk_size and current_chunk:
                    chunks.append(current_chunk)
                    
                    # Start new chunk with overlap
                    words = current_chunk.split()
                    overlap_words = min(len(words), overlap // 5)  # Approximate words in overlap
                    current_chunk = " ".join(words[-overlap_words:]) if overlap_words > 0 else ""
                    
                    # Add a separator if there's overlap
                    if current_chunk and not current_chunk.endswith("\n\n"):
                        current_chunk += "\n\n"
                
                # Add sentence chunk to current chunk
                if current_chunk and not current_chunk.endswith("\n\n"):
                    current_chunk += "\n\n"
                current_chunk += sentence_chunk
        else:
            # If adding this paragraph would exceed chunk_size, save the current chunk and start a new one
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk)
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_words = min(len(words), overlap // 5)  # Approximate words in overlap
                current_chunk = " ".join(words[-overlap_words:]) if overlap_words > 0 else ""
                
                # Add a separator if there's overlap
                if current_chunk and not current_chunk.endswith("\n\n"):
                    current_chunk += "\n\n"
            
            # Add paragraph to current chunk
            if current_chunk and not current_chunk.endswith("\n\n"):
                current_chunk += "\n\n"
            current_chunk += paragraph
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def chunk_by_sentence(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Chunk text by sentences."""
    import re
    
    # Split text into sentences
    # This is a simple regex for sentence splitting - could be improved
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If sentence itself is longer than chunk_size, split it by words
        if len(sentence) > chunk_size:
            words = sentence.split()
            current_sentence_chunk = ""
            
            for word in words:
                if len(current_sentence_chunk) + len(word) + 1 > chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = current_sentence_chunk
                    current_sentence_chunk = word
                else:
                    if current_sentence_chunk:
                        current_sentence_chunk += " "
                    current_sentence_chunk += word
            
            if current_sentence_chunk:
                if len(current_chunk) + len(current_sentence_chunk) > chunk_size:
                    chunks.append(current_chunk)
                    current_chunk = current_sentence_chunk
                else:
                    if current_chunk:
                        current_chunk += " "
                    current_chunk += current_sentence_chunk
        else:
            # If adding this sentence would exceed chunk_size, save the current chunk and start a new one
            if len(current_chunk) + len(sentence) + 1 > chunk_size and current_chunk:
                chunks.append(current_chunk)
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_words = min(len(words), overlap // 5)  # Approximate words in overlap
                current_chunk = " ".join(words[-overlap_words:]) if overlap_words > 0 else ""
            
            # Add sentence to current chunk
            if current_chunk:
                current_chunk += " "
            current_chunk += sentence
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def chunk_by_heading(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Chunk text by headings (for Markdown and RST)."""
    import re
    
    # Define regex patterns for headings
    md_heading_pattern = r'^#{1,6}\s+.+$'
    rst_heading_pattern = r'^[^\n]+\n[=\-~]+$'
    
    # Combine patterns
    heading_pattern = f"({md_heading_pattern}|{rst_heading_pattern})"
    
    # Split text by headings
    sections = re.split(f"(?m)^({heading_pattern})", text)
    
    # Recombine headings with their content
    sections_with_headings = []
    for i in range(0, len(sections), 2):
        if i + 1 < len(sections):
            sections_with_headings.append(sections[i] + sections[i+1])
        else:
            sections_with_headings.append(sections[i])
    
    # Process each section
    chunks = []
    for section in sections_with_headings:
        if len(section) <= chunk_size:
            chunks.append(section)
        else:
            # If section is too large, chunk it by paragraphs
            section_chunks = chunk_by_paragraph(section, chunk_size, overlap)
            chunks.extend(section_chunks)
    
    return chunks


def chunk_by_page_then_paragraph(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Chunk text by pages (for PDF) and then by paragraphs if needed."""
    # Split text by page markers
    pages = text.split("--- Page ")
    
    chunks = []
    for page in pages:
        if not page:
            continue
        
        # Extract page number if present
        page_parts = page.split(" ---", 1)
        if len(page_parts) > 1:
            page_num = page_parts[0]
            page_content = page_parts[1]
            page_header = f"--- Page {page_num} ---"
        else:
            page_header = ""
            page_content = page
        
        # If page is small enough, keep it as a single chunk
        if len(page_content) <= chunk_size:
            if page_header:
                chunks.append(f"{page_header}{page_content}")
            else:
                chunks.append(page_content)
        else:
            # Otherwise, chunk by paragraphs
            page_chunks = chunk_by_paragraph(page_content, chunk_size, overlap)
            
            # Add page header to first chunk
            if page_chunks and page_header:
                page_chunks[0] = f"{page_header}{page_chunks[0]}"
            
            chunks.extend(page_chunks)
    
    return chunks


async def create_document(
    title: str,
    description: Optional[str],
    doc_type: DocumentType,
    source: DocumentSource,
    user_id: str,
    file_path: Optional[str] = None,
    url: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
    db: Session = None,
) -> Tuple[bool, str]:
    """
    Create a new document.
    
    Args:
        title: Document title
        description: Document description
        doc_type: Document type
        source: Document source
        user_id: ID of the user creating the document
        file_path: Path to the document file (if applicable)
        url: URL of the document (if applicable)
        metadata: Additional metadata
        db: Database session
        
    Returns:
        A tuple of (success, document_id or error message)
    """
    try:
        # Create document
        document = Document(
            id=str(uuid.uuid4()),
            title=title,
            description=description,
            doc_type=doc_type,
            source=source,
            file_path=file_path,
            url=url,
            meta_data=metadata or {},
            created_by=user_id,
            status=DocumentStatus.PENDING,
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Start processing in the background
        # In a real implementation, this would be a background task
        # For now, we'll just process it directly
        asyncio.create_task(process_document(document.id, db))
        
        return True, document.id
    
    except Exception as e:
        return False, str(e)